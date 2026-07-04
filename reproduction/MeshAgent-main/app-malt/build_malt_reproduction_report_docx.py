from __future__ import annotations

import json
import math
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BASE = Path(__file__).resolve().parent
OUT = BASE / "MeshAgent_MALT_复现实验报告.docx"
CHART_DIR = BASE / "report_assets"
CHART_DIR.mkdir(exist_ok=True)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(96, 108, 124)
GRAY = RGBColor(242, 244, 247)
LIGHT_BLUE = RGBColor(232, 238, 245)
SOFT_YELLOW = RGBColor(255, 248, 232)
SOFT_RED = RGBColor(253, 237, 237)
GREEN = RGBColor(34, 121, 79)
RED = RGBColor(155, 28, 28)
GOLD = RGBColor(122, 90, 0)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)


def read_json(name: str) -> Any:
    return json.loads((BASE / name).read_text(encoding="utf-8"))


raw_summary = read_json("results_50_summary.json")
paper_full = read_json("results_malt_paper50_full_r3_summary.json")[0]
intent_summary = read_json("results_intentcheck50_v2_r3_summary.json")
llm_judge = read_json("results_llm_judge_hard_v1_summary.json")
intent_full = read_json("results_intentcheck50_v2_r3_Full.json")
intent_check = read_json("results_intentcheck50_v2_r3_FullIntentCheck.json")


def pct(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{value * 100:.2f}%"


def ratio(num: int, den: int) -> str:
    return f"{num}/{den}"


def difficulty(idx: int) -> str:
    if 1 <= idx <= 7 or 22 <= idx <= 45:
        return "easy"
    if 8 <= idx <= 14 or 46 <= idx <= 50:
        return "medium"
    if 15 <= idx <= 21:
        return "hard"
    return "other"


def compute_metrics(attempts: list[dict[str, Any]]) -> dict[str, Any]:
    total = len(attempts)
    raw_correct = sum(1 for item in attempts if item.get("correct"))
    answered = sum(1 for item in attempts if not item.get("abstained"))
    abstained = total - answered
    correct_answered = sum(1 for item in attempts if item.get("correct") and not item.get("abstained"))
    wrong_answered = sum(1 for item in attempts if (not item.get("correct")) and not item.get("abstained"))
    abstained_wrong = sum(1 for item in attempts if (not item.get("correct")) and item.get("abstained"))
    abstained_correct = sum(1 for item in attempts if item.get("correct") and item.get("abstained"))
    return {
        "total": total,
        "raw_correct": raw_correct,
        "answered": answered,
        "abstained": abstained,
        "correct_answered": correct_answered,
        "wrong_answered": wrong_answered,
        "abstained_wrong": abstained_wrong,
        "abstained_correct": abstained_correct,
        "raw_acc": raw_correct / total if total else None,
        "reliable": correct_answered / answered if answered else None,
        "abstain_rate": abstained / total if total else None,
    }


def metrics_by_difficulty(data: dict[str, Any]) -> dict[str, dict[str, Any]]:
    groups = {"easy": [], "medium": [], "hard": []}
    for query in data["queries"]:
        groups[difficulty(query["query_index"])].extend(query["attempts"])
    return {key: compute_metrics(value) for key, value in groups.items()}


difficulty_full = metrics_by_difficulty(intent_full)
difficulty_check = metrics_by_difficulty(intent_check)


def set_east_asia_font(run, font_name: str):
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None,
                 italic: bool | None = None, name: str = "Calibri", east_asia: str = "Microsoft YaHei"):
    run.font.name = name
    set_east_asia_font(run, east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6, line: float = 1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill: RGBColor):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), f"{fill.rgb:06X}")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_paragraph(doc: Document, text: str = "", style: str | None = None, size: float | None = None,
                  color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None,
                  align: WD_ALIGN_PARAGRAPH | None = None, before: float = 0, after: float = 6,
                  line: float = 1.10):
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p, before, after, line)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc: Document, text: str, level: int):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    if level == 1:
        set_paragraph_spacing(p, 16, 8, 1.10)
        size, color = 16, BLUE
    elif level == 2:
        set_paragraph_spacing(p, 12, 6, 1.10)
        size, color = 13, BLUE
    else:
        set_paragraph_spacing(p, 8, 4, 1.10)
        size, color = 12, DARK_BLUE
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, 0, 8, 1.167)
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, 0, 8, 1.167)
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_callout(doc: Document, label: str, text: str, fill: RGBColor = LIGHT_BLUE, accent: RGBColor = DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, 0, 2, 1.10)
    r1 = p.add_run(label + "  ")
    set_run_font(r1, 10.5, accent, True)
    r2 = p.add_run(text)
    set_run_font(r2, 10.5, BLACK)
    add_paragraph(doc, "", after=4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], note: str | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    set_table_geometry(table, widths)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], GRAY)
        p = header_cells[idx].paragraphs[0]
        set_paragraph_spacing(p, 0, 0, 1.05)
        r = p.add_run(header)
        set_run_font(r, 10, NAVY, True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            set_paragraph_spacing(p, 0, 0, 1.05)
            r = p.add_run(str(value))
            set_run_font(r, 9.5, BLACK)
            if idx > 0 and len(str(value)) <= 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_table_geometry(table, widths)
    if note:
        add_paragraph(doc, note, size=9, color=MUTED, italic=True, before=4, after=8)
    else:
        add_paragraph(doc, "", after=4)
    return table


def set_doc_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = BLUE
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = BLUE
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = DARK_BLUE
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def set_headers_footers(doc: Document):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, 0, 0, 1.0)
    r = p.add_run("MeshAgent MALT 复现实验报告")
    set_run_font(r, 9, MUTED, False)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, 0, 0, 1.0)
    r = p.add_run("第 ")
    set_run_font(r, 9, MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_end)
    r2 = p.add_run(" 页")
    set_run_font(r2, 9, MUTED)


def create_bar_chart(path: Path):
    width, height = 1100, 430
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 22)
    small = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 18)
    title = ImageFont.truetype("C:/Windows/Fonts/msyhbd.ttc", 28)
    draw.text((30, 25), "主要实验准确率对比", fill=(11, 37, 69), font=title)
    items = [
        ("Baseline", 0.32, (155, 28, 28)),
        ("+Constraints", 0.80, (46, 116, 181)),
        ("Simplified Full", 0.68, (122, 90, 0)),
        ("Paper-style Full raw", 0.8467, (34, 121, 79)),
        ("IntentCheck reliable", 1.00, (31, 77, 120)),
    ]
    left = 260
    top = 95
    bar_h = 38
    gap = 26
    max_w = 680
    for idx, (label, value, color) in enumerate(items):
        y = top + idx * (bar_h + gap)
        draw.text((35, y + 5), label, fill=(0, 0, 0), font=small)
        draw.rounded_rectangle((left, y, left + max_w, y + bar_h), radius=8, fill=(242, 244, 247))
        draw.rounded_rectangle((left, y, left + int(max_w * value), y + bar_h), radius=8, fill=color)
        draw.text((left + max_w + 18, y + 5), f"{value * 100:.2f}%", fill=color, font=font)
    draw.text((35, 390), "说明：IntentCheck reliable 只表示系统选择回答的子集全部正确，不代表 raw accuracy 达到 100%。", fill=(96, 108, 124), font=small)
    img.save(path)


def create_pipeline_chart(path: Path):
    width, height = 1100, 300
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title = ImageFont.truetype("C:/Windows/Fonts/msyhbd.ttc", 28)
    font = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 18)
    draw.text((30, 25), "复现实验与改进流程", fill=(11, 37, 69), font=title)
    boxes = [
        ("官方代码\nMALT", 40),
        ("约束检索\nCoT生成", 245),
        ("执行与\n结构检查", 450),
        ("置信度\n拒答", 655),
        ("IntentCheck\n意图后验验证", 860),
    ]
    y = 110
    for text, x in boxes:
        draw.rounded_rectangle((x, y, x + 150, y + 86), radius=12, fill=(232, 238, 245), outline=(46, 116, 181), width=2)
        lines = text.split("\n")
        for line_idx, line in enumerate(lines):
            draw.text((x + 24, y + 18 + line_idx * 28), line, fill=(11, 37, 69), font=font)
        if x < 860:
            draw.line((x + 160, y + 43, x + 195, y + 43), fill=(96, 108, 124), width=3)
            draw.polygon([(x + 195, y + 43), (x + 185, y + 36), (x + 185, y + 50)], fill=(96, 108, 124))
    draw.text((40, 230), "改进点：不改变 LLM 生成流程，在执行后补充 query-template-aware postcondition validation。", fill=(96, 108, 124), font=font)
    img.save(path)


def add_cover(doc: Document):
    add_paragraph(doc, "实验报告", size=12, color=BLUE, bold=True, after=6)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, 0, 4, 1.0)
    r = title.add_run("MeshAgent MALT 部分复现实验与改进分析")
    set_run_font(r, 24, NAVY, True)
    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, 0, 16, 1.10)
    r = subtitle.add_run("基于官方 GitHub 代码的近似复现、失败分析与 intent-level validation 改进")
    set_run_font(r, 13.5, MUTED)

    rows = [
        ["实验对象", "MeshAgent: Enabling Reliable Network Management with Large Language Models"],
        ["复现场景", "MALT / Network Lifecycle Management"],
        ["数据规模", "前 50 条 queries；主要实验每题 3 runs，共 150 attempts"],
        ["本地路径", r"F:\vs_program\meshagent_network\reproduction\MeshAgent-main\app-malt"],
        ["报告日期", "2026-06-30"],
    ]
    add_table(doc, ["项目", "说明"], rows, [1700, 7660])
    add_callout(
        doc,
        "核心结论",
        "本实验复现了 MeshAgent 在 MALT 场景中的核心运行流程，并基于 hard query 失败模式加入 template-level intent validation。改进版不提升 raw accuracy，但能显著减少高置信错答：同一 50 题 x 3 runs 中 wrong_answered 从 14 降至 0，代价是 abstain rate 从 13.33% 上升到 22.67%。",
        fill=SOFT_YELLOW,
        accent=GOLD,
    )
    doc.add_page_break()


def add_metric_cards(doc: Document):
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, [2340, 2340, 2340, 2340], indent=0)
    labels = ["Paper-style Full raw", "Paper-style Full reliable", "IntentCheck wrong answered", "IntentCheck abstain rate"]
    values = [pct(paper_full["raw_accuracy_before_abstention"]), pct(paper_full["reliable_accuracy"]), "14 -> 0", pct(intent_summary[1]["abstain_rate"])]
    colors = [BLUE, GREEN, RED, GOLD]
    for row_idx in range(2):
        for col_idx in range(4):
            cell = table.cell(row_idx, col_idx)
            set_cell_margins(cell, 120, 160, 120, 160)
            set_cell_shading(cell, LIGHT_BLUE if row_idx == 0 else RGBColor(255, 255, 255))
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, 0, 0, 1.0)
            text = labels[col_idx] if row_idx == 0 else values[col_idx]
            r = p.add_run(text)
            set_run_font(r, 9 if row_idx == 0 else 15, colors[col_idx] if row_idx == 1 else NAVY, True)
    add_paragraph(doc, "", after=6)


def build_doc():
    create_bar_chart(CHART_DIR / "accuracy_chart.png")
    create_pipeline_chart(CHART_DIR / "pipeline_chart.png")

    doc = Document()
    set_doc_styles(doc)
    set_headers_footers(doc)
    add_cover(doc)

    add_heading(doc, "1. 摘要", 1)
    add_metric_cards(doc)
    add_paragraph(
        doc,
        "本报告围绕论文 MeshAgent 在 MALT 场景中的复现展开。实验首先基于作者公开 GitHub 仓库运行 Baseline、+Constraints 与 Full MeshAgent 相关流程；随后补充更接近论文运行时机制的 paper-style Full runner，包括 query-specific constraints、tool retrieval、CoT 分步生成、执行错误修复、结构验证、confidence 与 abstention。最后，针对 hard query 中暴露出的 graph mutation 和 query intent 验证不足问题，本文实现了 template-level intent postcondition validation。",
    )
    add_paragraph(
        doc,
        "需要强调的是，本文的改进主要提升系统可靠性，而不是直接提升 LLM 原始代码生成能力。IntentCheck 将错误但会被输出的结果拦截下来，因此 reliable accuracy 提升明显；同时拒答率上升，说明该方法属于“少答但少错”的工程型改进。",
    )
    doc.add_picture(str(CHART_DIR / "accuracy_chart.png"), width=Inches(6.35))

    add_heading(doc, "2. 原文方法与复现范围", 1)
    add_heading(doc, "2.1 MeshAgent 的核心流程", 2)
    add_paragraph(
        doc,
        "MeshAgent 的核心思想不是单纯调整 prompt，而是将网络管理任务中的可复用 invariants 抽象为 constraints。运行时，系统为每个 query 检索相关 constraints 和 tools，将其加入 LLM 代码生成过程；执行后再用 validation tests 检查输出是否满足应用约束。若执行错误或约束违背出现，错误上下文会反馈给 LLM 进行 error reduction；若多轮修复后仍不可靠，则系统 abstain。",
    )
    doc.add_picture(str(CHART_DIR / "pipeline_chart.png"), width=Inches(6.35))
    add_heading(doc, "2.2 本地复现边界", 2)
    for item in [
        "使用官方公开仓库中的 app-malt 作为复现基础，但该仓库更像研究原型和模块化实验代码，不是完整一键复现实验包。",
        "本实验补充了 benchmark runner、每题 fresh graph 加载、执行超时、结果 JSON 记录、confidence 重分析与 intent validation 实验脚本。",
        "实验重点覆盖前 50 条 MALT queries，其中 easy 31 条、medium 12 条、hard 7 条；paper-style 结果采用每题 3 runs。",
        "未复现原文所有模型对比、fine-tuning、完整人工审核 constraint evolution 及作者内部运行环境。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. 实验设置", 1)
    add_table(
        doc,
        ["项目", "设置"],
        [
            ["运行环境", "Windows PowerShell，本地 Python 虚拟环境"],
            ["代码路径", r"F:\vs_program\meshagent_network\reproduction\MeshAgent-main\app-malt"],
            ["执行超时", "process_graph 默认 120 秒"],
            ["核心模型", "通过 .env 配置的 OpenAI-compatible chat model"],
            ["主实验规模", "50 queries x 3 runs = 150 attempts"],
            ["难度分布", "easy 31，medium 12，hard 7"],
        ],
        [1900, 7460],
    )
    add_heading(doc, "3.1 关键脚本", 2)
    add_table(
        doc,
        ["脚本", "作用"],
        [
            ["benchmark.py", "基础执行框架，修正 fresh graph、执行超时和返回类型记录"],
            ["run_reproduction_benchmark.py", "运行 Baseline、+Constraints、简化 Full 的 50 题 raw benchmark"],
            ["run_malt_paper_reproduction.py", "更接近论文流程的 Full MeshAgent runner"],
            ["query_intent_validators.py", "新增的 template-level intent postcondition validators"],
            ["run_hard_improvement_experiment.py", "对比 Full、Full+IntentCheck、Full+IntentDebug"],
            ["llm_intent_judge_reanalyze.py", "探索性 LLM-as-verifier 离线重分析脚本"],
        ],
        [2800, 6560],
    )

    add_heading(doc, "4. 复现实验结果", 1)
    add_heading(doc, "4.1 三组 raw benchmark", 2)
    raw_rows = [[r["experiment"], ratio(r["correct"], r["total"]), pct(r["accuracy"])] for r in raw_summary]
    add_table(doc, ["方法", "正确数", "Raw Accuracy"], raw_rows, [3300, 2600, 3460])
    add_paragraph(
        doc,
        "结果显示，显式 constraints 对 MALT 任务帮助明显。Baseline 只有 32.00%，加入全部 constraints 后达到 80.00%。简化 Full MeshAgent 未超过 +Constraints，说明 CoT 与初步 debug 若缺少稳定的运行时验证，可能破坏原本可由 constraints 直接解决的任务。",
    )
    add_heading(doc, "4.2 Paper-style Full MeshAgent", 2)
    add_table(
        doc,
        ["指标", "数值"],
        [
            ["Total attempts", str(paper_full["total_attempts"])],
            ["Raw correct", str(paper_full["raw_correct"])],
            ["Answered / Abstained", f'{paper_full["answered"]} / {paper_full["abstained"]}'],
            ["Correct answered / Wrong answered", f'{paper_full["correct_answered"]} / {paper_full["wrong_answered"]}'],
            ["Raw accuracy before abstention", pct(paper_full["raw_accuracy_before_abstention"])],
            ["Total accuracy", pct(paper_full["total_accuracy"])],
            ["Reliable accuracy", pct(paper_full["reliable_accuracy"])],
            ["Abstain rate", pct(paper_full["abstain_rate"])],
        ],
        [4200, 5160],
    )
    add_callout(
        doc,
        "解释",
        "Reliable accuracy = correct_answered / answered。它只衡量系统最终选择回答的子集是否可靠，不代表所有 attempts 都正确。因此 reliable accuracy 高并不等于 raw accuracy 高。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )

    add_heading(doc, "5. 失败分析", 1)
    add_paragraph(
        doc,
        "复现实验中，easy 和 medium 查询整体表现较好，主要错误集中在 hard graph manipulation 与复杂拓扑推理任务。典型失败包括新增节点挂错层级、边属性使用 relationship 而非 type、删除 switch 后仍保留目标节点、端口移除后容量不平衡、构造子图时缺少必要拓扑边，以及 redundancy/path-analysis 查询缺少对应验证规则。",
    )
    add_table(
        doc,
        ["失败类型", "代表 query", "影响"],
        [
            ["新增节点层级或边属性错误", "#02, #21", "输出图结构看似合法，但不满足命名层级或 RK_CONTAINS 语义"],
            ["删除/迁移操作未真实完成", "#16, #17", "返回 graph/list 通过基础类型检查，但没有满足 query intent"],
            ["子图构造缺边", "#15", "遗漏与目标节点相连的原始包含边"],
            ["未建模推理型查询", "#19", "冗余路径计算没有对应 postcondition verifier"],
        ],
        [2450, 1900, 5010],
    )

    add_heading(doc, "6. 改进实验：IntentCheck", 1)
    add_heading(doc, "6.1 方法说明", 2)
    add_paragraph(
        doc,
        "IntentCheck 不是 universal semantic verifier，而是 query-template-aware postcondition validation。它不针对具体 query id 写死答案，而是针对 MALT 中反复出现的任务类型设计通用检查函数：新增 packet switch、删除 packet switch、删除 ports 并保持容量平衡、构造包含指定类型节点的子图。",
    )
    for item in [
        "新增 packet_switch：检查节点类型、端口数量、端口容量、父节点层级和 RK_CONTAINS 边。",
        "删除 packet_switch：检查目标 switch 是否真的删除、图是否出现孤立节点、容量 spread 是否变差。",
        "删除 ports 并保持 balanced：检查每个 switch 的端口数量，并计算移除后容量 spread 是否达到可实现的最小值。",
        "构造子图：检查目标类型节点是否齐全，并保留原图中与目标节点相连的必要包含边。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6.2 50 queries x 3 runs 结果", 2)
    full = intent_summary[0]
    check = intent_summary[1]
    add_table(
        doc,
        ["指标", "Full", "Full+IntentCheck", "变化解释"],
        [
            ["Raw accuracy", pct(full["raw_accuracy_before_abstention"]), pct(check["raw_accuracy_before_abstention"]), "模型生成正确率不变"],
            ["Answered", str(full["answered"]), str(check["answered"]), "更保守，回答数量减少"],
            ["Wrong answered", str(full["wrong_answered"]), str(check["wrong_answered"]), "错误输出被拦截"],
            ["Abstained wrong", str(full["abstained_wrong"]), str(check["abstained_wrong"]), "成功拒答错误结果增加"],
            ["Abstained correct", str(full["abstained_correct"]), str(check["abstained_correct"]), "未额外误拦正确答案"],
            ["Reliable accuracy", pct(full["reliable_accuracy"]), pct(check["reliable_accuracy"]), "回答子集可靠性提升"],
            ["Abstain rate", pct(full["abstain_rate"]), pct(check["abstain_rate"]), "代价是拒答率上升"],
            ["Abstain recall", pct(full["abstain_recall"]), pct(check["abstain_recall"]), "错误结果召回更完整"],
        ],
        [2300, 1650, 2100, 3310],
    )
    add_callout(
        doc,
        "结论",
        "IntentCheck 使 wrong_answered 从 14 降至 0，reliable accuracy 从 89.23% 提升到 100.00%。但该提升并非模型 raw accuracy 提升，而是系统将不可靠结果拦截掉，因此应解释为可靠性提升，而非“准确率达到 100%”。",
        fill=SOFT_YELLOW,
        accent=GOLD,
    )

    add_heading(doc, "6.3 按难度分析", 2)
    rows = []
    for key, label in [("easy", "easy"), ("medium", "medium"), ("hard", "hard")]:
        f = difficulty_full[key]
        c = difficulty_check[key]
        rows.append([
            label,
            pct(f["raw_acc"]),
            pct(f["reliable"]),
            str(f["wrong_answered"]),
            pct(c["reliable"]),
            str(c["wrong_answered"]),
            pct(c["abstain_rate"]),
        ])
    add_table(
        doc,
        ["难度", "Full raw", "Full reliable", "Full wrong", "Intent reliable", "Intent wrong", "Intent abstain"],
        rows,
        [1050, 1350, 1500, 1300, 1600, 1350, 1210],
    )
    add_paragraph(
        doc,
        "hard 子集变化最明显：Full 的 hard reliable accuracy 只有 31.25%，wrong_answered 为 11；IntentCheck 后 hard wrong_answered 降为 0，但 hard abstain rate 达到 76.19%。这说明该方法有效降低高置信错答，但对 hard 问题的真实生成能力并未提升。",
    )

    add_heading(doc, "7. 探索性实验：LLM-as-verifier", 1)
    add_paragraph(
        doc,
        "为测试是否可以让模型“自己主动理解并检查”，本文新增 LLMIntentJudge 离线重分析脚本。该方法不看 ground truth，只基于 query、generated code、return preview 和 metadata 判断结果是否可信。",
    )
    add_table(
        doc,
        ["指标", "Hard 子集结果"],
        [
            ["Total attempts", str(llm_judge["total_attempts"])],
            ["Raw correct", str(llm_judge["raw_correct"])],
            ["Answered / Abstained", f'{llm_judge["answered"]} / {llm_judge["abstained"]}'],
            ["Correct answered / Wrong answered", f'{llm_judge["correct_answered"]} / {llm_judge["wrong_answered"]}'],
            ["Abstained correct", str(llm_judge["abstained_correct"])],
            ["Reliable accuracy", pct(llm_judge["reliable_accuracy"])],
            ["Abstain rate", pct(llm_judge["abstain_rate"])],
        ],
        [4200, 5160],
    )
    add_callout(
        doc,
        "观察",
        "LLM Judge 能主动怀疑 #19 redundancy query，但整体过于保守：hard 子集拒答率为 90.48%，并误拦 5 个正确结果，同时仍放过 2 个错答。因此它适合作为探索性分析，不适合作为最终主改进。",
        fill=SOFT_RED,
        accent=RED,
    )

    add_heading(doc, "8. 与原文约束机制的关系", 1)
    add_paragraph(
        doc,
        "原文的 constraints 也并非任意任务通用，而是 domain-specific / application-specific invariants。其泛化性主要体现在同一应用或相似任务中：query 很多，但底层结构规则有限。MeshAgent 通过 semi-automated constraint creation 和 engineer review 构建约束库，并在失败、低置信度或 abstention 后进行 dynamic constraint evolution。",
    )
    add_paragraph(
        doc,
        "本文的 IntentCheck 可理解为对原文 failure-driven constraint evolution 思路的简化模拟：在复现实验中观察 hard query 的失败模式后，将其抽象为若干可执行的 postcondition validators。它不针对具体 query id，但仍依赖人工设计的 intent templates，因此泛化性弱于原文完整的半自动约束生成与演化流程。",
    )

    add_heading(doc, "9. 局限性与未来工作", 1)
    for item in [
        "本实验只覆盖前 50 条 MALT queries，未完整复现全部 benchmark 和多模型对照。",
        "本地 GitHub 代码是官方公开 artifact，但不是完整一键复现实验包，部分 pipeline 需要自行补充。",
        "IntentCheck 的泛化性局限于已建模的 MALT graph mutation / topology construction intent templates。",
        "对于 #19 这类未建模的复杂推理型查询，需要新增 redundancy/path-analysis constraints 或程序化 verifier。",
        "LLM-as-verifier 能发现部分未建模错误，但目前误拒答较多，需结合更丰富的图事实摘要或形式化检查。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10. 总结", 1)
    add_paragraph(
        doc,
        "本文完成了 MeshAgent MALT 部分的近似复现，并验证了 constraints、query-specific retrieval、CoT、error reduction 与 confidence/abstention 对可靠性的作用。复现实验表明，paper-style Full MeshAgent 在前 50 条 queries、每题 3 runs 设置下达到 84.67% raw accuracy 和 92.13% reliable accuracy。",
    )
    add_paragraph(
        doc,
        "在此基础上，本文针对 hard query 失败模式实现了 query-template-aware intent validation。该方法将 wrong_answered 从 14/150 降至 0/150，说明其能有效降低高置信错误输出；但 raw accuracy 保持不变，abstain rate 上升，说明该改进本质是提升系统可靠性，而不是提升模型生成能力。整体而言，该实验完成了“论文阅读理解 - 代码运行复现 - 失败分析与优化尝试”的完整闭环。",
    )

    add_heading(doc, "附录 A：关键运行命令", 1)
    commands = [
        ("三组 raw benchmark", r"..\venv\Scripts\python.exe run_reproduction_benchmark.py --limit 50 --timeout 120 --output-prefix results_50"),
        ("Paper-style Full", r'..\venv\Scripts\python.exe run_malt_paper_reproduction.py --limit 50 --runs 3 --timeout 120 --groups "Full MeshAgent" --output-prefix results_malt_paper50_full_r3'),
        ("IntentCheck 50 题", r"..\venv\Scripts\python.exe run_hard_improvement_experiment.py --query-indices $q --runs 3 --timeout 120 --intent-debug-loops 0 --output-prefix results_intentcheck50_v2_r3"),
        ("LLM Judge hard 子集", r"..\venv\Scripts\python.exe llm_intent_judge_reanalyze.py --input results_intentcheck50_v2_r3_Full.json --query-indices 15,16,17,18,19,20,21 --output results_llm_judge_hard_v1.json --summary-output results_llm_judge_hard_v1_summary.json"),
    ]
    add_table(doc, ["用途", "命令"], commands, [2300, 7060])

    add_heading(doc, "附录 B：输出文件索引", 1)
    add_table(
        doc,
        ["文件", "说明"],
        [
            ["results_50_summary.json", "Baseline / +Constraints / simplified Full 的 50 题 raw 结果"],
            ["results_malt_paper50_full_r3_summary.json", "Paper-style Full MeshAgent 主复现结果"],
            ["results_intentcheck50_v2_r3_summary.json", "IntentCheck 改进实验总表"],
            ["results_llm_judge_hard_v1_summary.json", "LLM-as-verifier 探索实验 hard 子集结果"],
            ["query_intent_validators.py", "IntentCheck 规则实现"],
            ["llm_intent_judge_reanalyze.py", "LLM judge 离线重分析脚本"],
        ],
        [3900, 5460],
    )

    # Preset audit markers are encoded through styles, explicit page setup, and table DXA geometry.
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
