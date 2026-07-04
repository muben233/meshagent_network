from __future__ import annotations

import json
import math
import re
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BASE = Path(__file__).resolve().parent
PROJECT_ROOT = BASE.parents[2]
TEMPLATE = next(PROJECT_ROOT.glob("北京邮电大学研究生学位论文模板*.docx"))
OUT = BASE / "MeshAgent_MALT_复现实验报告_模板版.docx"
ASSET_DIR = BASE / "report_assets_template"
ASSET_DIR.mkdir(exist_ok=True)

BLACK = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(64, 64, 64)
MID_GRAY = RGBColor(105, 105, 105)
TABLE_HEADER = RGBColor(230, 236, 242)
TABLE_SUBTLE = RGBColor(246, 248, 250)
LIGHT_BLUE = RGBColor(221, 235, 247)
SOFT_RED = RGBColor(252, 228, 214)
SOFT_GREEN = RGBColor(226, 239, 218)
ACCENT_BLUE = RGBColor(37, 99, 145)
ACCENT_GREEN = RGBColor(68, 132, 89)
ACCENT_RED = RGBColor(171, 71, 66)
ACCENT_YELLOW = RGBColor(196, 145, 2)
WHITE = RGBColor(255, 255, 255)

BODY_EA = "宋体"
HEADING_EA = "黑体"
COVER_EA = "黑体"
LATIN = "Times New Roman"


def read_json(name: str) -> Any:
    return json.loads((BASE / name).read_text(encoding="utf-8"))


raw_summary = read_json("results_50_summary.json")
paper_full = read_json("results_malt_paper50_full_r3_summary.json")[0]
intent_summary = read_json("results_intentcheck50_v2_r3_summary.json")
llm_judge = read_json("results_llm_judge_hard_v1_summary.json")
intent_full = read_json("results_intentcheck50_v2_r3_Full.json")
intent_check = read_json("results_intentcheck50_v2_r3_FullIntentCheck.json")


def pct(value: float | None, digits: int = 2) -> str:
    if value is None:
        return "-"
    return f"{value * 100:.{digits}f}%"


def ratio(num: int, den: int) -> str:
    return f"{num}/{den}"


def difficulty(query_index: int) -> str:
    if 1 <= query_index <= 7 or 22 <= query_index <= 45:
        return "easy"
    if 8 <= query_index <= 14 or 46 <= query_index <= 50:
        return "medium"
    if 15 <= query_index <= 21:
        return "hard"
    return "other"


def compute_metrics(attempts: list[dict[str, Any]]) -> dict[str, Any]:
    total = len(attempts)
    raw_correct = sum(1 for item in attempts if item.get("correct"))
    answered = sum(1 for item in attempts if not item.get("abstained"))
    abstained = total - answered
    correct_answered = sum(
        1 for item in attempts if item.get("correct") and not item.get("abstained")
    )
    wrong_answered = sum(
        1 for item in attempts if (not item.get("correct")) and not item.get("abstained")
    )
    abstained_wrong = sum(
        1 for item in attempts if (not item.get("correct")) and item.get("abstained")
    )
    abstained_correct = sum(
        1 for item in attempts if item.get("correct") and item.get("abstained")
    )
    return {
        "total": total,
        "raw_correct": raw_correct,
        "answered": answered,
        "abstained": abstained,
        "correct_answered": correct_answered,
        "wrong_answered": wrong_answered,
        "abstained_wrong": abstained_wrong,
        "abstained_correct": abstained_correct,
        "raw_acc": raw_correct / total if total else None,
        "total_acc": correct_answered / total if total else None,
        "reliable": correct_answered / answered if answered else None,
        "abstain_rate": abstained / total if total else None,
    }


def metrics_by_difficulty(data: dict[str, Any]) -> dict[str, dict[str, Any]]:
    grouped: dict[str, list[dict[str, Any]]] = {"easy": [], "medium": [], "hard": []}
    for query in data["queries"]:
        grouped[difficulty(query["query_index"])].extend(query["attempts"])
    return {name: compute_metrics(items) for name, items in grouped.items()}


difficulty_full = metrics_by_difficulty(intent_full)
difficulty_check = metrics_by_difficulty(intent_check)


def rgb_hex(color: RGBColor) -> str:
    return str(color)


def set_east_asia_font(run_or_style: Any, font_name: str) -> None:
    if hasattr(run_or_style, "_element"):
        element = run_or_style._element
    else:
        element = run_or_style.element
    rpr = element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), font_name)


def set_run_font(
    run: Any,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    east_asia: str = BODY_EA,
    latin: str = LATIN,
) -> None:
    run.font.name = latin
    set_east_asia_font(run, east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def clear_paragraph(paragraph: Any) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    section.different_first_page_header_footer = True

    for style_name in ["Normal", "Body Text"]:
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = LATIN
            set_east_asia_font(style, BODY_EA)
            style.font.size = Pt(12)
            style.font.color.rgb = BLACK

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = LATIN
            set_east_asia_font(style, HEADING_EA)
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = BLACK

    header = section.header
    for paragraph in header.paragraphs:
        clear_paragraph(paragraph)
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run("MeshAgent MALT 复现实验报告")
    set_run_font(run, 9, MID_GRAY, east_asia=BODY_EA)

    footer = section.footer
    for paragraph in footer.paragraphs:
        clear_paragraph(paragraph)
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)

    for container in [section.first_page_header, section.first_page_footer]:
        for paragraph in container.paragraphs:
            clear_paragraph(paragraph)


def add_page_number(paragraph: Any) -> None:
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, MID_GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for item in [fld_begin, instr, fld_sep, text, fld_end]:
        run._r.append(item)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, MID_GRAY)


def set_paragraph_spacing(
    paragraph: Any,
    before: float = 0,
    after: float = 6,
    line: float = 1.25,
    first_line: float | None = None,
) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if first_line is not None:
        pf.first_line_indent = Pt(first_line)


def add_para(
    doc: Document,
    text: str = "",
    *,
    size: float = 12,
    bold: bool = False,
    color: RGBColor = BLACK,
    align: WD_ALIGN_PARAGRAPH | None = None,
    before: float = 0,
    after: float = 6,
    line: float = 1.25,
    first_line: float | None = None,
    east_asia: str = BODY_EA,
) -> Any:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before, after, line, first_line)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size, color, bold=bold, east_asia=east_asia)
    return p


def add_mixed_para(
    doc: Document,
    parts: Iterable[tuple[str, bool]],
    *,
    size: float = 12,
    after: float = 6,
    first_line: float | None = 24,
) -> Any:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, after, 1.25, first_line)
    for text, bold in parts:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return p


def add_chapter(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 14, 12, 1.20)
    run = p.add_run(title)
    set_run_font(run, 16, BLACK, bold=True, east_asia=HEADING_EA)


def add_section_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, 10, 6, 1.20)
    run = p.add_run(title)
    set_run_font(run, 14, BLACK, bold=True, east_asia=HEADING_EA)


def add_subsection_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 8, 4, 1.20)
    run = p.add_run(title)
    set_run_font(run, 13, BLACK, bold=True, east_asia=HEADING_EA)


def add_caption(doc: Document, text: str) -> None:
    add_para(
        doc,
        text,
        size=10.5,
        color=DARK_GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=3,
        after=8,
        line=1.10,
    )


def add_bullet(doc: Document, text: str, *, level: int = 0) -> Any:
    p = doc.add_paragraph()
    apply_bullet_numbering(doc, p, level)
    set_paragraph_spacing(p, 0, 3, 1.18)
    run = p.add_run(text)
    set_run_font(run, 12)
    return p


def ensure_bullet_numbering(doc: Document) -> int:
    cached = getattr(doc, "_meshagent_bullet_num_id", None)
    if cached is not None:
        return cached

    numbering = doc.part.numbering_part.element
    abstract_nums = numbering.findall(qn("w:abstractNum"))
    nums = numbering.findall(qn("w:num"))
    max_abstract = max(
        [int(item.get(qn("w:abstractNumId"))) for item in abstract_nums] or [0]
    )
    max_num = max([int(item.get(qn("w:numId"))) for item in nums] or [0])
    abstract_id = max_abstract + 1
    num_id = max_num + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "hybridMultilevel")
    abstract.append(multi)

    for ilvl, left, hanging in [(0, 420, 220), (1, 760, 220)]:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet")
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "·")
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "space")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        ppr.append(ind)
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), BODY_EA)
        rfonts.set(qn("w:hAnsi"), BODY_EA)
        rfonts.set(qn("w:eastAsia"), BODY_EA)
        rpr.append(rfonts)
        for node in [start, num_fmt, lvl_text, suffix, lvl_jc, ppr, rpr]:
            lvl.append(node)
        abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    setattr(doc, "_meshagent_bullet_num_id", num_id)
    return num_id


def apply_bullet_numbering(doc: Document, paragraph: Any, level: int = 0) -> None:
    num_id = ensure_bullet_numbering(doc)
    ppr = paragraph._p.get_or_add_pPr()
    existing = ppr.find(qn("w:numPr"))
    if existing is not None:
        ppr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)


def add_note_box(doc: Document, label: str, text: str, fill: RGBColor = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [8000])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, 0, 0, 1.20)
    r1 = p.add_run(label)
    set_run_font(r1, 11.5, BLACK, bold=True, east_asia=HEADING_EA)
    r2 = p.add_run(text)
    set_run_font(r2, 11.5, BLACK)
    add_para(doc, "", after=4)


def set_cell_shading(cell: Any, fill: RGBColor) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), rgb_hex(fill))


def set_cell_margins(cell: Any, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell: Any, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table: Any, widths: list[int], indent: int = 0) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def format_cell(
    cell: Any,
    text: str,
    *,
    size: float = 10.5,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    fill: RGBColor | None = None,
    color: RGBColor = BLACK,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_spacing(p, 0, 0, 1.10)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)
    if fill is not None:
        set_cell_shading(cell, fill)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    *,
    font_size: float = 10.5,
    first_col_left: bool = False,
) -> Any:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        format_cell(table.rows[0].cells[i], header, size=font_size, bold=True, fill=TABLE_HEADER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if first_col_left and i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            fill = TABLE_SUBTLE if len(table.rows) % 2 == 1 else None
            format_cell(cells[i], value, size=font_size, align=align, fill=fill)
    add_para(doc, "", after=4)
    return table


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    chunks = re.split(r"(\s+)", text)
    lines: list[str] = []
    current = ""
    for chunk in chunks:
        candidate = current + chunk
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if width <= max_width or not current:
            current = candidate
        else:
            lines.append(current.strip())
            current = chunk
    if current.strip():
        lines.append(current.strip())
    return lines


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for name in candidates:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_bar_chart(path: Path) -> None:
    data = [
        ("Baseline", 0.32, ACCENT_RED),
        ("+Constraints", 0.80, ACCENT_GREEN),
        ("Paper-style Full\nraw", paper_full["raw_accuracy_before_abstention"], ACCENT_BLUE),
        ("Paper-style Full\ntotal", paper_full["total_accuracy"], ACCENT_YELLOW),
        ("IntentCheck\ntotal", intent_summary[1]["total_accuracy"], DARK_GRAY),
        ("IntentCheck\nreliable", intent_summary[1]["reliable_accuracy"], ACCENT_BLUE),
    ]
    w, h = 1500, 760
    image = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, True)
    label_font = load_font(24)
    small_font = load_font(22)
    axis_font = load_font(20)
    draw.text((60, 40), "MALT 前 50 题复现实验主要指标", fill=(0, 0, 0), font=title_font)
    draw.text(
        (60, 88),
        "注：reliable accuracy 只统计未拒答样本，不能与 raw / total accuracy 直接等价比较。",
        fill=(90, 90, 90),
        font=axis_font,
    )
    left, top, bottom, right = 110, 160, 620, 1440
    draw.line((left, bottom, right, bottom), fill=(70, 70, 70), width=2)
    draw.line((left, top, left, bottom), fill=(70, 70, 70), width=2)
    for tick in range(0, 101, 20):
        y = bottom - int((bottom - top) * tick / 100)
        draw.line((left - 8, y, right, y), fill=(225, 225, 225), width=1)
        draw.text((50, y - 12), f"{tick}%", fill=(80, 80, 80), font=axis_font)
    gap = 34
    bar_w = (right - left - gap * (len(data) + 1)) // len(data)
    for idx, (label, value, color) in enumerate(data):
        x0 = left + gap + idx * (bar_w + gap)
        bar_h = int((bottom - top) * value)
        y0 = bottom - bar_h
        draw.rounded_rectangle((x0, y0, x0 + bar_w, bottom), radius=8, fill=tuple(color))
        pct_text = f"{value * 100:.1f}%"
        tw = draw.textbbox((0, 0), pct_text, font=label_font)[2]
        draw.text((x0 + (bar_w - tw) / 2, y0 - 35), pct_text, fill=(0, 0, 0), font=label_font)
        label_lines = label.split("\n")
        for line_no, line in enumerate(label_lines):
            tw = draw.textbbox((0, 0), line, font=small_font)[2]
            draw.text(
                (x0 + (bar_w - tw) / 2, bottom + 20 + line_no * 28),
                line,
                fill=(50, 50, 50),
                font=small_font,
            )
    image.save(path)


def draw_pipeline_chart(path: Path) -> None:
    w, h = 1500, 520
    image = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, True)
    box_font = load_font(23, True)
    small_font = load_font(19)
    draw.text((60, 35), "复现流程与改进位置", fill=(0, 0, 0), font=title_font)
    boxes = [
        ("Question", "MALT 查询"),
        ("RAG", "检索约束与工具"),
        ("CoT", "生成图操作代码"),
        ("Exec", "执行 process_graph"),
        ("MyChecker", "答案/拒答判定"),
        ("IntentCheck", "后置意图校验"),
    ]
    colors = [
        (230, 236, 242),
        (221, 235, 247),
        (226, 239, 218),
        (242, 242, 242),
        (255, 242, 204),
        (252, 228, 214),
    ]
    x, y, bw, bh, gap = 55, 155, 205, 120, 30
    for i, ((title, desc), color) in enumerate(zip(boxes, colors)):
        x0 = x + i * (bw + gap)
        draw.rounded_rectangle((x0, y, x0 + bw, y + bh), radius=14, fill=color, outline=(150, 150, 150), width=2)
        tw = draw.textbbox((0, 0), title, font=box_font)[2]
        draw.text((x0 + (bw - tw) / 2, y + 26), title, fill=(0, 0, 0), font=box_font)
        tw2 = draw.textbbox((0, 0), desc, font=small_font)[2]
        draw.text((x0 + (bw - tw2) / 2, y + 70), desc, fill=(70, 70, 70), font=small_font)
        if i < len(boxes) - 1:
            ax0 = x0 + bw + 5
            ax1 = x0 + bw + gap - 5
            ay = y + bh // 2
            draw.line((ax0, ay, ax1, ay), fill=(80, 80, 80), width=3)
            draw.polygon([(ax1, ay), (ax1 - 12, ay - 8), (ax1 - 12, ay + 8)], fill=(80, 80, 80))
    note = "IntentCheck 是本次复现后的改进实验：它不提高 raw correctness，但减少错误回答并提高可靠回答比例。"
    for i, line in enumerate(wrap_text(draw, note, small_font, 1300)):
        draw.text((85, 345 + i * 28), line, fill=(70, 70, 70), font=small_font)
    image.save(path)


def build_cover(doc: Document) -> None:
    add_para(doc, "", after=30)
    add_para(
        doc,
        "课程论文报告",
        size=22,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=4,
        after=26,
        east_asia=COVER_EA,
    )
    title = add_para(
        doc,
        "",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=36,
        line=1.25,
        east_asia=COVER_EA,
    )
    run = title.add_run("MeshAgent 在 MALT 网络管理任务上的")
    set_run_font(run, 18, BLACK, bold=True, east_asia=COVER_EA)
    run.add_break(WD_BREAK.LINE)
    run = title.add_run("复现实验报告")
    set_run_font(run, 18, BLACK, bold=True, east_asia=COVER_EA)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [2200, 5200])
    info = [
        ("论文名称", "MeshAgent 在 MALT 网络管理任务上的复现实验报告"),
        ("姓    名", "（待填写）"),
        ("学    号", "（待填写）"),
        ("学 院(系)", "（待填写）"),
        ("专    业", "（待填写）"),
        ("指导教师", "（待填写）"),
    ]
    for row, (left, right) in zip(table.rows, info):
        format_cell(row.cells[0], left, size=12, bold=True, fill=TABLE_SUBTLE)
        format_cell(row.cells[1], right, size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, "", after=48)
    add_para(
        doc,
        "2026年6月30日",
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=12,
        after=0,
    )
    doc.add_page_break()


def build_abstract(doc: Document) -> None:
    add_para(
        doc,
        "摘  要",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=8,
        after=14,
        east_asia=HEADING_EA,
    )
    add_para(
        doc,
        "本文围绕 MeshAgent: Enabling Reliable Network Management with Large Language Models "
        "中的 MALT 应用展开复现。复现过程包括论文阅读理解、代码运行复现和面向 hard "
        "问题的改进尝试三个阶段。实验首先对 Baseline、加入约束的模型以及 Full MeshAgent "
        "进行前 50 题对比；随后按照更接近原文的设置，对 Full MeshAgent 执行 50 个查询、"
        "每题 3 次的实验，并统计 raw accuracy、total accuracy、reliable accuracy 和拒答率。"
        "在此基础上，本文进一步设计了 IntentCheck 后置意图校验实验，用于降低模型在复杂推理题上"
        "的错误回答风险。",
        size=12,
        first_line=24,
        after=8,
    )
    add_para(
        doc,
        "复现实验表明，MALT 任务中显式约束对准确率提升明显；更接近原文的 Full MeshAgent "
        "在前 50 题三次运行中 raw accuracy 达到 84.67%，total accuracy 为 78.00%，"
        "reliable accuracy 为 92.13%。改进实验中，IntentCheck 未提升 total accuracy，"
        "但将错误回答数从 14 降为 0，使 reliable accuracy 从 89.23% 提升到 100.00%，"
        "代价是拒答率从 13.33% 增至 22.67%。这说明当前改进更偏向可靠性过滤，而不是直接增强"
        "模型解决 hard 查询的能力。",
        size=12,
        first_line=24,
        after=10,
    )
    add_para(
        doc,
        "关键词：MeshAgent；MALT；网络管理；大语言模型；约束检索；复现实验；可靠性校验",
        size=12,
        bold=False,
        after=12,
    )
    doc.add_page_break()


def build_toc_like_page(doc: Document) -> None:
    add_para(
        doc,
        "目  录",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=8,
        after=18,
        east_asia=HEADING_EA,
    )
    toc_items = [
        "第一章  实验背景与复现目标",
        "第二章  实验设置与复现流程",
        "第三章  复现实验结果",
        "第四章  失败分析与改进实验",
        "第五章  结论与后续工作",
        "附录 A  运行命令与结果文件",
    ]
    for item in toc_items:
        add_para(doc, item, size=12, before=0, after=8, line=1.25)
    doc.add_page_break()


def build_main_content(doc: Document) -> None:
    chart1 = ASSET_DIR / "accuracy_chart.png"
    chart2 = ASSET_DIR / "pipeline_chart.png"
    draw_bar_chart(chart1)
    draw_pipeline_chart(chart2)

    add_chapter(doc, "第一章  实验背景与复现目标")
    add_section_title(doc, "1.1 论文方法理解")
    add_para(
        doc,
        "MeshAgent 的核心目标是让大语言模型在网络管理任务中更可靠地生成、执行并检查代码。"
        "论文并不是单纯让模型直接回答问题，而是将领域约束、工具使用、代码执行和自检过程组合成"
        "一个 agent 流程。MALT 应用中的查询主要围绕网络拓扑图进行，包括节点查询、属性更新、"
        "带宽计算、包含关系统计以及新增网络实体等任务。",
        first_line=24,
    )
    add_para(
        doc,
        "本次复现将重点放在 MALT 应用，原因是该应用的数据结构明确、答案可通过图操作验证，"
        "同时包含 easy、medium 和 hard 三类任务，适合观察约束、执行反馈和拒答机制对最终结果的影响。",
        first_line=24,
    )
    add_section_title(doc, "1.2 复现目标")
    add_bullet(doc, "跑通论文 GitHub 项目中 MALT 应用的主要流程，确认 Baseline、+Constraints 和 Full MeshAgent 的差异。")
    add_bullet(doc, "修正本地复现中图状态复用、超时控制、代码执行命名空间等会影响结果可信度的问题。")
    add_bullet(doc, "采用更接近原文的 Full MeshAgent 设置，对前 50 个查询执行每题 3 次的实验。")
    add_bullet(doc, "在复现基础上尝试通用性较强的后置校验改进，并重点分析 hard 查询中的失败模式。")

    doc.add_page_break()
    add_chapter(doc, "第二章  实验设置与复现流程")
    add_section_title(doc, "2.1 数据与查询范围")
    easy_n = 31
    medium_n = 12
    hard_n = 7
    add_para(
        doc,
        f"实验选择 MALT 数据集中前 50 个查询，其中 easy {easy_n} 题、medium {medium_n} 题、"
        f"hard {hard_n} 题。为了接近原文对不确定性的处理，Full MeshAgent 主实验对每个查询运行 3 次，"
        "共形成 150 次尝试。除原始正确率外，报告还统计模型未拒答部分的可靠准确率。",
        first_line=24,
    )
    add_table(
        doc,
        ["难度", "题目数量", "查询编号范围"],
        [
            ["easy", "31", "1-7，22-45"],
            ["medium", "12", "8-14，46-50"],
            ["hard", "7", "15-21"],
        ],
        [1800, 1800, 4400],
    )
    add_caption(doc, "表 2-1  前 50 个 MALT 查询的难度分布")

    add_section_title(doc, "2.2 流程还原")
    add_para(
        doc,
        "复现流程保留了论文中 Full MeshAgent 的关键环节：基于查询检索约束和工具说明，生成代码，"
        "在本地图状态上执行代码，再通过 checker 计算答案正确性和拒答情况。与早期 smoke benchmark "
        "相比，后续脚本为每一道题重新加载初始图状态，避免前一题的图修改污染后续查询。",
        first_line=24,
    )
    doc.add_picture(str(chart2), width=Inches(6.2))
    add_caption(doc, "图 2-1  MALT 复现流程与本次改进插入位置")

    doc.add_page_break()
    add_section_title(doc, "2.3 指标说明")
    add_table(
        doc,
        ["指标", "含义", "解释重点"],
        [
            ["raw accuracy", "不考虑拒答前的原始正确率", "反映模型生成结果本身是否正确"],
            ["total accuracy", "correct_answered / total_attempts", "把拒答也计入总尝试，更接近端到端可用性"],
            ["reliable accuracy", "correct_answered / answered", "只看模型选择回答的样本是否可靠"],
            ["abstain rate", "abstained / total_attempts", "越高说明模型越保守"],
        ],
        [1900, 3100, 3000],
        first_col_left=True,
    )
    add_caption(doc, "表 2-2  本报告使用的主要评价指标")
    add_note_box(
        doc,
        "指标提醒：",
        " reliable accuracy 不是整体正确率。它只在模型选择回答的样本上统计，因此 100% 可能意味着“回答更谨慎”，"
        "并不意味着所有题都已经解决。",
        fill=SOFT_RED,
    )

    doc.add_page_break()
    add_chapter(doc, "第三章  复现实验结果")
    add_section_title(doc, "3.1 前 50 题三组对比")
    rows = []
    for item in raw_summary:
        rows.append(
            [
                item["experiment"],
                ratio(item["correct"], item["total"]),
                pct(item["accuracy"]),
                {
                    "Baseline": "无约束，主要作为下限参考",
                    "+Constraints": "显式约束后提升明显",
                    "Full MeshAgent": "早期简化版 Full，结果低于约束组",
                }.get(item["experiment"], ""),
            ]
        )
    add_table(
        doc,
        ["实验组", "正确数", "准确率", "说明"],
        rows,
        [1900, 1500, 1500, 3100],
        first_col_left=True,
    )
    add_caption(doc, "表 3-1  前 50 题三组对比结果")
    add_para(
        doc,
        "结果显示，约束信息对 MALT 任务非常关键。Baseline 只有 32.00%，加入全量约束后达到 80.00%。"
        "早期 Full MeshAgent 只有 68.00%，这说明本地复现中的 Full 流程如果没有严格还原原文的多次运行、"
        "拒答和自检统计方式，可能会低估方法能力。",
        first_line=24,
    )

    add_section_title(doc, "3.2 更接近原文设置的 Full MeshAgent")
    add_table(
        doc,
        ["统计项", "数值"],
        [
            ["查询数 / 每题运行次数", f'{paper_full["total_queries"]} / {paper_full["runs_per_query"]}'],
            ["总尝试次数", str(paper_full["total_attempts"])],
            ["raw correct", ratio(paper_full["raw_correct"], paper_full["total_attempts"])],
            ["answered / abstained", f'{paper_full["answered"]} / {paper_full["abstained"]}'],
            ["correct_answered / wrong_answered", f'{paper_full["correct_answered"]} / {paper_full["wrong_answered"]}'],
            ["raw accuracy", pct(paper_full["raw_accuracy_before_abstention"])],
            ["total accuracy", pct(paper_full["total_accuracy"])],
            ["reliable accuracy", pct(paper_full["reliable_accuracy"])],
            ["abstain rate", pct(paper_full["abstain_rate"])],
        ],
        [3600, 4400],
        first_col_left=True,
    )
    add_caption(doc, "表 3-2  Paper-style Full MeshAgent 前 50 题三次运行结果")
    doc.add_picture(str(chart1), width=Inches(6.25))
    add_caption(doc, "图 3-1  主要实验指标对比")
    add_para(
        doc,
        "在更接近原文的设置下，Full MeshAgent 的 raw accuracy 达到 84.67%，明显高于早期简化 Full 的 68.00%。"
        "但 total accuracy 为 78.00%，说明拒答和错误回答仍然会影响端到端可用性。"
        "从可靠性角度看，92.13% 的 reliable accuracy 表明模型在选择回答时多数是可信的，但仍存在少量错误回答。",
        first_line=24,
    )

    add_section_title(doc, "3.3 分难度表现")
    rows = []
    names = {"easy": "easy", "medium": "medium", "hard": "hard"}
    for key in ["easy", "medium", "hard"]:
        item = difficulty_full[key]
        rows.append(
            [
                names[key],
                str(item["total"] // 3),
                ratio(item["raw_correct"], item["total"]),
                pct(item["raw_acc"]),
                f'{item["correct_answered"]}/{item["answered"]}',
                pct(item["reliable"]),
                str(item["wrong_answered"]),
            ]
        )
    add_table(
        doc,
        ["难度", "题数", "raw correct", "raw acc.", "可靠回答", "reliable acc.", "错误回答"],
        rows,
        [1000, 900, 1500, 1200, 1300, 1300, 1100],
        font_size=9.5,
    )
    add_caption(doc, "表 3-3  Full MeshAgent 在不同难度上的表现")
    add_para(
        doc,
        "easy 和 medium 查询整体表现较好，hard 查询明显成为瓶颈。前 50 题中 hard 只有 7 题，"
        "但 21 次尝试中 raw correct 仅 5 次，且错误回答集中出现在这一组。这说明后续优化不应只看总体准确率，"
        "而应单独分析 hard 查询的推理链、图遍历范围和答案格式约束。",
        first_line=24,
    )

    doc.add_page_break()
    add_chapter(doc, "第四章  失败分析与改进实验")
    add_section_title(doc, "4.1 主要失败类型")
    add_bullet(doc, "实体创建类失败：新增 packet_switch 等任务中容易出现 None 节点或包含关系未补齐。")
    add_bullet(doc, "聚合计算类失败：按 AGG_BLOCK、RACK 或 CHASSIS 汇总时，容易漏掉层级范围或单位换算。")
    add_bullet(doc, "排序/Top-k 类失败：hard 查询中常见先过滤范围再排序的多步推理错误。")
    add_bullet(doc, "拒答策略不稳定：模型有时对正确答案拒答，有时又对错误代码给出回答。")
    add_para(
        doc,
        "这些失败不是单一 API 调用错误造成的，而是图结构语义、约束检索、代码生成和结果验证共同作用的结果。"
        "其中 hard 查询最容易暴露“代码能运行但意图不一致”的问题，例如查询要求在指定子图中找最大容量节点，"
        "模型却可能在更大范围内排序，导致答案格式看似正确但语义错误。",
        first_line=24,
    )

    add_section_title(doc, "4.2 IntentCheck 改进实验")
    full = intent_summary[0]
    check = intent_summary[1]
    add_table(
        doc,
        ["实验组", "raw acc.", "total acc.", "reliable acc.", "拒答率", "错误回答", "正确拒答"],
        [
            [
                "Full",
                pct(full["raw_accuracy_before_abstention"]),
                pct(full["total_accuracy"]),
                pct(full["reliable_accuracy"]),
                pct(full["abstain_rate"]),
                str(full["wrong_answered"]),
                str(full["abstained_wrong"]),
            ],
            [
                "Full+IntentCheck",
                pct(check["raw_accuracy_before_abstention"]),
                pct(check["total_accuracy"]),
                pct(check["reliable_accuracy"]),
                pct(check["abstain_rate"]),
                str(check["wrong_answered"]),
                str(check["abstained_wrong"]),
            ],
        ],
        [1900, 1200, 1200, 1500, 1100, 1100, 1100],
        font_size=9.5,
        first_col_left=True,
    )
    add_caption(doc, "表 4-1  IntentCheck 对 Full MeshAgent 的影响")
    add_para(
        doc,
        "IntentCheck 的作用不是重新求解题目，而是在模型给出答案后，根据查询语义生成并执行一组后置条件检查。"
        "从结果看，它没有改变 raw accuracy，也没有提高 total accuracy；但它把错误回答从 14 次降为 0 次，"
        "将 reliable accuracy 提升到 100.00%。这说明该改进更适合作为可靠性防线，而不是作为提升 hard 题求解能力的核心方法。",
        first_line=24,
    )

    doc.add_page_break()
    rows = []
    for key in ["easy", "medium", "hard"]:
        base = difficulty_full[key]
        improved = difficulty_check[key]
        rows.append(
            [
                key,
                pct(base["raw_acc"]),
                pct(base["reliable"]),
                str(base["wrong_answered"]),
                pct(improved["reliable"]),
                str(improved["wrong_answered"]),
                pct(improved["abstain_rate"]),
            ]
        )
    add_table(
        doc,
        ["难度", "Full raw", "Full reliable", "Full 错答", "Check reliable", "Check 错答", "Check 拒答率"],
        rows,
        [900, 1200, 1300, 1100, 1500, 1100, 1300],
        font_size=9,
    )
    add_caption(doc, "表 4-2  IntentCheck 分难度结果")
    add_note_box(
        doc,
        "结果解释：",
        " hard 组 reliable accuracy 变为 100.00%，主要因为错误 hard 尝试被拒答过滤；同时 hard 拒答率达到 76.19%，"
        "所以不能把该结果解释为 hard 题已经被真正解决。",
        fill=SOFT_GREEN,
    )

    add_section_title(doc, "4.3 LLM-as-verifier 探索")
    add_table(
        doc,
        ["统计项", "结果"],
        [
            ["实验范围", "hard 查询 7 题，共 21 次尝试"],
            ["raw correct", ratio(llm_judge["raw_correct"], llm_judge["total_attempts"])],
            ["answered / abstained", f'{llm_judge["answered"]} / {llm_judge["abstained"]}'],
            ["correct_answered / wrong_answered", f'{llm_judge["correct_answered"]} / {llm_judge["wrong_answered"]}'],
            ["reliable accuracy", pct(llm_judge["reliable_accuracy"])],
            ["abstain rate", pct(llm_judge["abstain_rate"])],
        ],
        [3600, 4400],
        first_col_left=True,
    )
    add_caption(doc, "表 4-3  LLM-as-verifier 在 hard 子集上的探索结果")
    add_para(
        doc,
        "LLM-as-verifier 的探索结果并不理想：它对 hard 子集过于保守，拒答率达到 90.48%，"
        "同时仍放过了 2 次错误回答。该实验说明，单纯再调用一个大模型进行意图判断并不一定更稳定；"
        "如果没有明确的可执行后置条件，验证器本身也会受到语义理解偏差影响。",
        first_line=24,
    )

    doc.add_page_break()
    add_chapter(doc, "第五章  结论与后续工作")
    add_section_title(doc, "5.1 复现结论")
    add_para(
        doc,
        "本文已完成论文阅读理解、代码运行复现和初步优化改进三个阶段。从复现角度看，当前结果已经能够支持"
        "“约束检索和 agent 式执行反馈能显著提升 MALT 网络管理任务可靠性”的主要结论。前 50 题三次运行的"
        "Paper-style Full MeshAgent 结果与早期 smoke benchmark 相比更稳定，也更接近原文实验思路。",
        first_line=24,
    )
    add_para(
        doc,
        "从改进角度看，IntentCheck 证明了后置语义校验可以降低错误回答风险，但它更像是可靠性过滤器，"
        "而不是通用求解器。若后续继续提高 accuracy，重点应放在 hard 查询的主动理解、子图范围推断、"
        "多步图遍历代码生成和可执行验证条件自动生成上。",
        first_line=24,
    )
    add_section_title(doc, "5.2 后续优化方向")
    add_bullet(doc, "面向 hard 查询建立通用的查询意图解析层，将范围、目标实体、聚合方式、排序规则和单位要求结构化。")
    add_bullet(doc, "把后置检查从人工整理规则扩展为由模型生成候选检查、再由程序验证检查是否可执行的机制。")
    add_bullet(doc, "增加错误案例库，让系统从错答和拒答中动态更新约束，这与原文中约束不断演化的思想一致。")
    add_bullet(doc, "单独报告 hard 子集准确率，避免总体指标被 easy/medium 查询掩盖。")

    doc.add_page_break()
    add_chapter(doc, "附录 A  运行命令与结果文件")
    add_section_title(doc, "A.1 关键运行命令")
    commands = [
        "python run_malt_paper_reproduction.py --query-indices 1..50 --runs 3 --timeout 120 --groups full",
        "python run_hard_improvement_experiment.py --query-indices 1..50 --runs 3 --timeout 120 --intent-debug-loops 0 --output-prefix results_intentcheck50_v2_r3",
        "python llm_intent_judge_reanalyze.py --input results_intentcheck50_v2_r3_Full.json --output-prefix results_llm_judge_hard_v1",
    ]
    for cmd in commands:
        p = add_para(
            doc,
            cmd,
            size=10,
            before=0,
            after=5,
            line=1.10,
            first_line=None,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )
        for run in p.runs:
            run.font.name = "Consolas"
            set_east_asia_font(run, "Microsoft YaHei")

    add_section_title(doc, "A.2 主要结果文件")
    add_table(
        doc,
        ["文件", "用途"],
        [
            ["results_50_summary.json", "前 50 题 Baseline / +Constraints / Full 三组摘要"],
            ["results_malt_paper50_full_r3_summary.json", "更接近原文设置的 Full MeshAgent 摘要"],
            ["results_intentcheck50_v2_r3_summary.json", "IntentCheck 改进实验摘要"],
            ["results_llm_judge_hard_v1_summary.json", "LLM-as-verifier hard 子集探索摘要"],
        ],
        [3600, 4400],
        font_size=10,
        first_col_left=True,
    )
    add_caption(doc, "表 A-1  本报告引用的主要结果文件")


def main() -> None:
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    configure_document(doc)
    build_cover(doc)
    build_abstract(doc)
    build_toc_like_page(doc)
    build_main_content(doc)
    doc.save(str(OUT))
    print(f"saved: {OUT}")


if __name__ == "__main__":
    main()
